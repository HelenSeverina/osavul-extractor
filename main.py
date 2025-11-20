import os
import sys
import csv
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dateutil import parser as dateparser
from dateutil import tz
from datetime import datetime

def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color_elem = OxmlElement("w:color")
    color_elem.set(qn("w:val"), color)
    rPr.append(color_elem)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def try_parse_datetime(value: str):
    if not value or not isinstance(value, str):
        return None, False
    val = value.strip()

    try:
        dt = dateparser.parse(val)
        if dt is None:
            return None, False
        if dt.tzinfo is None:
            kyiv = tz.gettz("Europe/Kyiv")
            dt = dt.replace(tzinfo=kyiv)
        kyiv = tz.gettz("Europe/Kyiv")
        dt_kyiv = dt.astimezone(kyiv)
        has_time = "T" in val or ":" in val
        return dt_kyiv, has_time
    except Exception as e:
        print(f"Error parsing datetime: {e}")
        return None, False

def format_date_time(dt: datetime, include_time: bool):
    if dt is None:
        return None, None
    date_str = dt.strftime("%d.%m.%Y")
    time_str = dt.strftime("%H:%M") if include_time else None
    return date_str, time_str

PLATFORM_MAP_UPDATED = {
    "TELEGRAM": "на Telegram-каналі",
    "YOUTUBE": "на Youtube-каналі",
    "WEB": "на Веб-сторінці",
    "FACEBOOK": "на Facebook-сторінці",
    "TWITTER": "у соціальній мережі Х",
    "VK": "у соціальній мережі VK",
    "TIKTOK": "на TikTok-сторінці"
}

def parse_records_from_csv(path: str):
    records = []
    with open(path, newline='', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            records.append({
                "date": row.get("date", ""),
                "url": row.get("url", ""),
                "platform": row.get("platform", ""),
                "source_name": row.get("source_name", "")
            })
    return records

def build_output_lines_updated(records):
    out_lines = []
    for rec in records:
        raw_date = rec.get("date", "")
        url = rec.get("url", "")
        platform_raw = rec.get("platform", "")
        source_name = rec.get("source_name", "")

        plat_key = platform_raw.strip()
        plat_label = PLATFORM_MAP_UPDATED.get(plat_key, plat_key)

        dt, has_time = try_parse_datetime(raw_date)
        date_str, time_str = format_date_time(dt, has_time) if dt else (None, None)

        if date_str and time_str and plat_label and source_name and url:
            line = f"{date_str} о {time_str} {plat_label} \"{source_name}\" за посиланням {url}"
        elif date_str and plat_label and source_name and url:
            line = f"{date_str} {plat_label} \"{source_name}\" за посиланням {url}"
        else:
            line = " ".join([
                date_str or "",
                f"о {time_str}" if time_str else "",
                plat_label or "",
                f"\"{source_name}\"" if source_name else "",
                f"за посиланням {url}" if url else ""
            ]).strip()

        out_lines.append((line, url))
    return out_lines

def write_output_docx(lines, out_path: str):
    doc = Document()
    for ln, url in lines:
        if url:
            before, _, after = ln.partition(url)
            p = doc.add_paragraph(before.strip() + " ")
            add_hyperlink(p, url.strip(), url.strip())
            if after.strip():
                p.add_run(" " + after.strip())
        else:
            doc.add_paragraph(ln)
    doc.save(out_path)

def main():
    if getattr(sys, 'frozen', False):
        base_dir = Path(sys.executable).resolve().parent
    else:
        base_dir = Path(__file__).resolve().parent

    input_path = base_dir / "input.csv"
    output_path = base_dir / "output.docx"

    if not input_path.exists():
        print("❌ Файл input.csv не знайдено.")
        print("🔷 Покладіть input.csv у ту саму папку, що і виконуваний файл.")
        if os.name == "nt":
            input("Натисніть Enter для виходу...")
        return

    try:
        records = parse_records_from_csv(str(input_path))
        lines_generated_updated = build_output_lines_updated(records)
        write_output_docx(lines_generated_updated, str(output_path))
        print(f"✅ Готово! Файл збережено як {output_path.name}")
    except Exception as e:
        print(f"⚠️ Помилка: {e}")
        if os.name == "nt":
            input("Натисніть Enter для виходу...")

if __name__ == "__main__":
    main()
